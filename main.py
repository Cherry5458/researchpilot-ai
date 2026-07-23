# ─────────────────────────────────────────────────────────────────────────────
# main.py  –  FastAPI Application Entry Point  (v4 – LangGraph Edition)
# ─────────────────────────────────────────────────────────────────────────────

import io
import logging
import urllib.parse
from datetime import date

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, JSONResponse
from pydantic import BaseModel, field_validator
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import ALLOWED_ORIGINS
from services.langgraph_orchestrator import run_langgraph_pipeline
from services.topic_suggestions import get_trending_topics
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-8s | %(name)s – %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)

# ── In-memory cache ───────────────────────────────────────────────────────────
_last_result: dict = {}

app = FastAPI(title="ResearchPilot AI", version="4.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.mount("/static", StaticFiles(directory="frontend"), name="static")

@app.get("/")
async def serve_frontend():
    return FileResponse("frontend/index.html")


class PaperRequest(BaseModel):
    topic: str

    @field_validator("topic")
    @classmethod
    def topic_not_empty(cls, v: str) -> str:
        v = v.strip()
        if not v:
            raise ValueError("Topic cannot be empty.")
        if len(v) < 5:
            raise ValueError("Topic must be at least 5 characters.")
        return v


def _build_docx(topic, paper_text, references, keywords) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(topic); r.bold = True; r.font.size = Pt(18)
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run(f"ResearchPilot AI  |  {date.today().strftime('%B %d, %Y')}").font.size = Pt(10)
    kp = doc.add_paragraph()
    kp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kp.add_run(f"Keywords: {', '.join(keywords)}").font.size = Pt(10)
    doc.add_paragraph()
    for line in paper_text.splitlines():
        s = line.strip()
        if s.startswith("## "):
            h = doc.add_heading(s[3:].strip(), level=1)
            h.runs[0].font.size = Pt(13)
        elif s:
            doc.add_paragraph(s).style.font.size = Pt(12)
    doc.add_heading("References", level=1)
    for ref in references:
        doc.add_paragraph(ref, style="List Number")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.get("/health")
async def health():
    return {"status": "ok", "service": "ResearchPilot AI v4 – LangGraph Edition"}


@app.get("/trending-topics")
async def trending_topics():
    """Return 6 trending research topic suggestions."""
    try:
        topics = await get_trending_topics()
        return JSONResponse({"topics": topics})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/generate-paper")
async def generate_paper(request: PaperRequest):
    """Run the full 9-node LangGraph pipeline and return a PDF."""
    global _last_result
    logger.info("POST /generate-paper | topic=%s", request.topic)
    try:
        pdf_bytes, metadata = await run_langgraph_pipeline(request.topic)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc))
    except Exception as exc:
        logger.exception("Pipeline error: %s", exc)
        raise HTTPException(status_code=500, detail=f"Pipeline failed: {str(exc)}")

    _last_result = {**metadata, "pdf_bytes": pdf_bytes}
    safe_name = urllib.parse.quote(metadata["validated_topic"][:60].replace(" ", "_"))

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition":   f'attachment; filename="{safe_name}_ResearchPaper.pdf"',
            "X-Word-Count":          str(metadata["word_count"]),
            "X-Reference-Count":     str(metadata["reference_count"]),
            "X-Validated-Topic":     metadata["validated_topic"],
            "X-Originality-Score":   str(metadata.get("plagiarism", {}).get("originality_score", "N/A")),
        },
    )


@app.get("/download-docx")
async def download_docx():
    if not _last_result:
        raise HTTPException(status_code=404, detail="No paper generated yet.")
    docx_bytes = _build_docx(
        _last_result["validated_topic"], _last_result["paper_text"],
        _last_result["references"], _last_result["keywords"],
    )
    safe_name = urllib.parse.quote(_last_result["validated_topic"][:60].replace(" ", "_"))
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_ResearchPaper.docx"'},
    )


@app.get("/preview")
async def preview():
    if not _last_result:
        raise HTTPException(status_code=404, detail="No paper generated yet.")
    return JSONResponse({
        "topic":            _last_result["validated_topic"],
        "keywords":         _last_result["keywords"],
        "word_count":       _last_result["word_count"],
        "reference_count":  _last_result["reference_count"],
        "paper_text":       _last_result["paper_text"],
        "references":       _last_result["references"],
        "plagiarism":       _last_result.get("plagiarism", {}),
    })


@app.get("/plagiarism-report")
async def plagiarism_report():
    """Return the plagiarism/originality report for the last generated paper."""
    if not _last_result:
        raise HTTPException(status_code=404, detail="No paper generated yet.")
    p = _last_result.get("plagiarism", {})
    if not p:
        raise HTTPException(status_code=404, detail="No plagiarism report available.")
    return JSONResponse(p)
